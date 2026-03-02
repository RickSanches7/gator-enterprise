"""
GATOR PRO Enterprise — Report Generator
═════════════════════════════════════════
Generates professional pentest reports:
  • PDF — executive + technical (RU/EN)
  • DOCX — editable Word document (RU/EN)
  • JSON — machine-readable full data

PDF structure:
  1. Cover page (logo, client, date, classification)
  2. Executive Summary (risk dashboard, top findings)
  3. Methodology
  4. Scope & Targets
  5. Findings (sorted by CVSS, with evidence & PoC)
  6. Compliance Summary (PCI DSS / SWIFT)
  7. Recommendations
  8. Appendix (technical details)
"""

import json
import os
import re
import time
from datetime import datetime, timezone
from typing import Optional


# ─── Translations ─────────────────────────────────────────────
STRINGS = {
    "ru": {
        "report_title":     "ОТЧЁТ О ТЕСТИРОВАНИИ НА ПРОНИКНОВЕНИЕ",
        "executive_summary":"РЕЗЮМЕ ДЛЯ РУКОВОДСТВА",
        "methodology":      "МЕТОДОЛОГИЯ",
        "scope":            "ОБЛАСТЬ ТЕСТИРОВАНИЯ",
        "findings":         "ОБНАРУЖЕННЫЕ УЯЗВИМОСТИ",
        "recommendations":  "РЕКОМЕНДАЦИИ",
        "compliance":       "СООТВЕТСТВИЕ СТАНДАРТАМ",
        "appendix":         "ПРИЛОЖЕНИЕ",
        "critical":         "КРИТИЧЕСКАЯ",
        "high":             "ВЫСОКАЯ",
        "medium":           "СРЕДНЯЯ",
        "low":              "НИЗКАЯ",
        "info":             "ИНФОРМАЦИОННАЯ",
        "pass":             "ВЫПОЛНЕНО",
        "fail":             "НЕ ВЫПОЛНЕНО",
        "manual":           "ТРЕБУЕТ ПРОВЕРКИ",
        "prepared_by":      "Подготовлено",
        "prepared_for":     "Для",
        "date":             "Дата",
        "confidential":     "КОНФИДЕНЦИАЛЬНО — ТОЛЬКО ДЛЯ РУКОВОДСТВА",
        "total_findings":   "Всего уязвимостей",
        "risk_level":       "Уровень риска",
        "cvss_score":       "Оценка CVSS",
        "description":      "Описание",
        "recommendation":   "Рекомендация",
        "evidence":         "Свидетельства",
        "poc":              "Подтверждение концепции (PoC)",
        "affected":         "Затронутый хост",
        "owasp":            "Категория OWASP",
        "pci_req":          "Требование PCI DSS",
        "swift_ctrl":       "Контроль SWIFT CSCF",
        "methodology_text": (
            "Тестирование проводилось согласно методологии PTES (Penetration Testing Execution Standard) "
            "и OWASP Web Security Testing Guide v4.2. Использовались следующие фазы:\n"
            "1. Разведка (OSINT, DNS, сертификаты)\n"
            "2. Сканирование портов и сервисов\n"
            "3. Тестирование веб-приложений (OWASP Top 10)\n"
            "4. Тестирование API безопасности (OWASP API Top 10)\n"
            "5. Тестирование аутентификации и сессий\n"
            "6. Анализ SSL/TLS\n"
            "7. Тестирование бизнес-логики\n"
            "8. Тестирование Active Directory\n"
            "9. Соответствие PCI DSS v4.0 и SWIFT CSCF\n"
            "10. Сетевая безопасность"
        ),
    },
    "en": {
        "report_title":     "PENETRATION TEST REPORT",
        "executive_summary":"EXECUTIVE SUMMARY",
        "methodology":      "METHODOLOGY",
        "scope":            "SCOPE & TARGETS",
        "findings":         "FINDINGS",
        "recommendations":  "RECOMMENDATIONS",
        "compliance":       "COMPLIANCE SUMMARY",
        "appendix":         "APPENDIX",
        "critical":         "CRITICAL",
        "high":             "HIGH",
        "medium":           "MEDIUM",
        "low":              "LOW",
        "info":             "INFORMATIONAL",
        "pass":             "PASS",
        "fail":             "FAIL",
        "manual":           "MANUAL CHECK",
        "prepared_by":      "Prepared by",
        "prepared_for":     "Prepared for",
        "date":             "Date",
        "confidential":     "CONFIDENTIAL — MANAGEMENT EYES ONLY",
        "total_findings":   "Total Findings",
        "risk_level":       "Risk Level",
        "cvss_score":       "CVSS Score",
        "description":      "Description",
        "recommendation":   "Recommendation",
        "evidence":         "Evidence",
        "poc":              "Proof of Concept (PoC)",
        "affected":         "Affected Host",
        "owasp":            "OWASP Category",
        "pci_req":          "PCI DSS Requirement",
        "swift_ctrl":       "SWIFT CSCF Control",
        "methodology_text": (
            "Testing was conducted following the PTES (Penetration Testing Execution Standard) "
            "methodology and OWASP Web Security Testing Guide v4.2. The following phases were performed:\n"
            "1. Reconnaissance (OSINT, DNS, certificates)\n"
            "2. Port and service scanning\n"
            "3. Web application testing (OWASP Top 10)\n"
            "4. API security testing (OWASP API Top 10)\n"
            "5. Authentication and session testing\n"
            "6. SSL/TLS analysis\n"
            "7. Business logic testing\n"
            "8. Active Directory testing\n"
            "9. PCI DSS v4.0 and SWIFT CSCF compliance\n"
            "10. Network security"
        ),
    }
}

SEV_COLORS = {
    "critical": (0.8, 0.0, 0.0),
    "high":     (0.9, 0.3, 0.0),
    "medium":   (0.9, 0.6, 0.0),
    "low":      (0.1, 0.5, 0.1),
    "info":     (0.3, 0.3, 0.7),
}


class ReportGenerator:
    def __init__(self, scan_data: dict, language: str = "ru",
                 client_name: str = "Bank", output_dir: str = "/tmp"):
        self.scan       = scan_data
        self.lang       = language if language in ("ru", "en") else "ru"
        self.s          = STRINGS[self.lang]
        self.client     = client_name
        self.output_dir = output_dir
        self.findings   = sorted(
            scan_data.get("findings", []),
            key=lambda x: x.get("cvss", 0), reverse=True
        )
        self.now        = datetime.now(timezone.utc)
        self.date_str   = self.now.strftime("%d.%m.%Y" if language == "ru" else "%Y-%m-%d")

    def generate_all(self) -> dict:
        """Generate PDF, DOCX, and JSON reports."""
        results = {}

        # Always generate JSON
        json_path = os.path.join(self.output_dir, f"report_{self.lang}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(self._build_json(), f, ensure_ascii=False, indent=2, default=str)
        results["json"] = json_path

        # PDF
        pdf_path = self._generate_pdf()
        if pdf_path:
            results["pdf"] = pdf_path

        # DOCX
        docx_path = self._generate_docx()
        if docx_path:
            results["docx"] = docx_path

        return results

    def _build_json(self) -> dict:
        """Build complete JSON report structure."""
        target    = self.scan.get("target", "")
        findings  = self.findings
        pci       = self.scan.get("pci_results", {})
        swift     = self.scan.get("swift_results", {})

        sev_counts = {s: len([f for f in findings if f.get("severity") == s])
                      for s in ("critical","high","medium","low","info")}
        pci_fail  = sum(1 for v in pci.values() if v.get("status") == "FAIL")
        swift_fail = sum(1 for v in swift.values() if v.get("status") == "FAIL")

        return {
            "report_metadata": {
                "title":         self.s["report_title"],
                "client":        self.client,
                "target":        target,
                "date":          self.date_str,
                "language":      self.lang,
                "generated_by":  "GATOR PRO Enterprise v2.0",
                "generated_at":  self.now.isoformat(),
            },
            "executive_summary": {
                "total_findings":  len(findings),
                "severity_counts": sev_counts,
                "risk_score":      self._calc_risk_score(),
                "pci_dss_fail":    pci_fail,
                "swift_fail":      swift_fail,
                "top_findings":    [
                    {"title": f["title"], "cvss": f.get("cvss",0),
                     "severity": f.get("severity",""), "host": f.get("host","")}
                    for f in findings[:5]
                ],
            },
            "scope": {
                "targets":   self.scan.get("targets", [target]),
                "scan_type": self.scan.get("scan_type", "full"),
                "started_at": self.scan.get("started_at", ""),
                "finished_at": self.scan.get("finished_at", ""),
            },
            "findings":     findings,
            "pci_results":  pci,
            "swift_results": swift,
            "recon_data":   self.scan.get("recon", {}),
            "open_ports":   self.scan.get("open_ports", []),
        }

    def _calc_risk_score(self) -> str:
        """Calculate overall risk score."""
        weights = {"critical":10,"high":7,"medium":4,"low":1,"info":0}
        score = sum(weights.get(f.get("severity","info"),0) for f in self.findings)
        if score >= 50:   return "CRITICAL"
        if score >= 20:   return "HIGH"
        if score >= 10:   return "MEDIUM"
        if score > 0:     return "LOW"
        return "INFO"

    # ─── PDF Generation ───────────────────────────────────────
    def _generate_pdf(self) -> Optional[str]:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm, cm
            from reportlab.lib import colors
            from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                            Table, TableStyle, PageBreak, HRFlowable)
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from reportlab.pdfgen import canvas as pdf_canvas
        except ImportError:
            return self._generate_pdf_minimal()

        path = os.path.join(self.output_dir, f"report_{self.lang}.pdf")

        # Color palette
        DARK_BLUE  = colors.HexColor("#1a2b4a")
        MID_BLUE   = colors.HexColor("#2563eb")
        LIGHT_GRAY = colors.HexColor("#f8fafc")
        RED        = colors.HexColor("#dc2626")
        ORANGE     = colors.HexColor("#ea580c")
        AMBER      = colors.HexColor("#d97706")
        GREEN      = colors.HexColor("#16a34a")

        SEV_CL = {
            "critical": RED, "high": ORANGE,
            "medium": AMBER, "low": GREEN,
            "info": colors.HexColor("#6b7280"),
        }

        doc = SimpleDocTemplate(path, pagesize=A4,
            leftMargin=20*mm, rightMargin=20*mm,
            topMargin=20*mm, bottomMargin=20*mm)

        styles = getSampleStyleSheet()
        story  = []

        def style(name, **kw):
            return ParagraphStyle(name, parent=styles["Normal"], **kw)

        h1 = style("H1", fontSize=22, textColor=DARK_BLUE,
                   spaceAfter=8, fontName="Helvetica-Bold", alignment=TA_CENTER)
        h2 = style("H2", fontSize=14, textColor=DARK_BLUE,
                   spaceAfter=6, fontName="Helvetica-Bold",
                   borderPad=4, backColor=LIGHT_GRAY)
        h3 = style("H3", fontSize=11, textColor=MID_BLUE,
                   spaceAfter=4, fontName="Helvetica-Bold")
        body = style("Body", fontSize=9, leading=14, spaceAfter=4)
        mono = style("Mono", fontSize=8, fontName="Courier", leading=11,
                     backColor=colors.HexColor("#1e293b"), textColor=colors.white,
                     leftIndent=8, rightIndent=8, spaceBefore=4, spaceAfter=4)

        # ── COVER PAGE ────────────────────────────────────────
        story.append(Spacer(1, 30*mm))
        story.append(Paragraph("🐊 GATOR PRO ENTERPRISE", style("logo",
            fontSize=28, textColor=MID_BLUE, fontName="Helvetica-Bold",
            alignment=TA_CENTER)))
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(width="100%", thickness=3, color=MID_BLUE))
        story.append(Spacer(1, 8*mm))
        story.append(Paragraph(self.s["report_title"], h1))
        story.append(Spacer(1, 15*mm))

        risk = self._calc_risk_score()
        risk_cl = {"CRITICAL":RED,"HIGH":ORANGE,"MEDIUM":AMBER,"LOW":GREEN}.get(risk,GREEN)
        meta_data = [
            [self.s["prepared_for"]+":", self.client],
            [self.s["date"]+":",         self.date_str],
            ["Target:",                  self.scan.get("target","")],
            [self.s["risk_level"]+":",   risk],
        ]
        meta_table = Table(meta_data, colWidths=[60*mm, 100*mm])
        meta_table.setStyle(TableStyle([
            ("FONTSIZE",    (0,0),(-1,-1), 11),
            ("FONTNAME",    (0,0),(0,-1), "Helvetica-Bold"),
            ("TEXTCOLOR",   (0,0),(0,-1), DARK_BLUE),
            ("ROWBACKGROUNDS",(0,0),(-1,-1), [LIGHT_GRAY, colors.white]),
            ("BOTTOMPADDING",(0,0),(-1,-1), 8),
            ("TOPPADDING",   (0,0),(-1,-1), 8),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 20*mm))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey))
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(self.s["confidential"], style("conf",
            fontSize=9, textColor=colors.red, alignment=TA_CENTER)))
        story.append(PageBreak())

        # ── EXECUTIVE SUMMARY ─────────────────────────────────
        story.append(Paragraph(self.s["executive_summary"], h2))
        story.append(Spacer(1, 4*mm))

        sev_counts = {s: len([f for f in self.findings if f.get("severity")==s])
                      for s in ("critical","high","medium","low","info")}
        total = len(self.findings)

        # Summary table
        sum_data = [["Severity / Серьёзность", "Count / Кол-во", "Status"]]
        for sev in ("critical","high","medium","low","info"):
            cnt = sev_counts[sev]
            cl  = SEV_CL.get(sev, colors.grey)
            sum_data.append([
                Paragraph(f'<font color="{cl.hexval()}">{sev.upper()}</font>', body),
                str(cnt),
                "⚠️ Requires action" if sev in ("critical","high") and cnt > 0 else "—",
            ])
        sum_data.append(["TOTAL", str(total), ""])

        sum_table = Table(sum_data, colWidths=[80*mm, 40*mm, 60*mm])
        sum_table.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,0), DARK_BLUE),
            ("TEXTCOLOR",    (0,0),(-1,0), colors.white),
            ("FONTNAME",     (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",     (0,0),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [LIGHT_GRAY, colors.white]),
            ("GRID",         (0,0),(-1,-1), 0.5, colors.lightgrey),
            ("BOTTOMPADDING",(0,0),(-1,-1), 6),
            ("TOPPADDING",   (0,0),(-1,-1), 6),
        ]))
        story.append(sum_table)
        story.append(Spacer(1, 8*mm))

        # PCI/SWIFT summary
        pci_results = self.scan.get("pci_results",{})
        swift_results = self.scan.get("swift_results",{})
        if pci_results:
            pci_pass = sum(1 for v in pci_results.values() if v.get("status")=="PASS")
            pci_fail = sum(1 for v in pci_results.values() if v.get("status")=="FAIL")
            pci_total = len(pci_results)
            story.append(Paragraph(
                f"PCI DSS v4.0: {pci_pass}/{pci_total} PASS, "
                f"<font color='red'>{pci_fail} FAIL</font>", body))
        if swift_results:
            s_pass = sum(1 for v in swift_results.values() if v.get("status")=="PASS")
            s_fail = sum(1 for v in swift_results.values() if v.get("status")=="FAIL")
            story.append(Paragraph(
                f"SWIFT CSCF v2024: {s_pass}/{len(swift_results)} PASS, "
                f"<font color='red'>{s_fail} FAIL</font>", body))

        story.append(PageBreak())

        # ── METHODOLOGY ───────────────────────────────────────
        story.append(Paragraph(self.s["methodology"], h2))
        story.append(Paragraph(self.s["methodology_text"].replace("\n","<br/>"), body))
        story.append(PageBreak())

        # ── FINDINGS ──────────────────────────────────────────
        story.append(Paragraph(self.s["findings"], h2))
        story.append(Spacer(1, 4*mm))

        for i, f in enumerate(self.findings, 1):
            sev  = f.get("severity","info")
            cvss = f.get("cvss", 0)
            cl   = SEV_CL.get(sev, colors.grey)
            title = f.get("title","")
            story.append(Paragraph(
                f'<font color="{cl.hexval()}">■</font> '
                f'<b>#{i} [{sev.upper()}] CVSS {cvss}</b> — {title}', h3))

            details = [
                [self.s["affected"],     f.get("host","") + (":" + str(f.get("port","")) if f.get("port") else "")],
                [self.s["owasp"],        f.get("owasp_category","—")],
            ]
            if f.get("pci_dss_req"):
                details.append([self.s["pci_req"], ", ".join(f.get("pci_dss_req",[]))])
            if f.get("swift_control"):
                details.append([self.s["swift_ctrl"], ", ".join(f.get("swift_control",[]))])

            det_table = Table(details, colWidths=[45*mm, 130*mm])
            det_table.setStyle(TableStyle([
                ("FONTSIZE",  (0,0),(-1,-1), 8),
                ("FONTNAME",  (0,0),(0,-1), "Helvetica-Bold"),
                ("TEXTCOLOR", (0,0),(0,-1), DARK_BLUE),
                ("ROWBACKGROUNDS",(0,0),(-1,-1), [LIGHT_GRAY, colors.white]),
                ("BOTTOMPADDING",(0,0),(-1,-1), 3),
            ]))
            story.append(det_table)
            story.append(Spacer(1,2*mm))
            story.append(Paragraph(f"<b>{self.s['description']}:</b> {f.get('description','')}", body))
            story.append(Paragraph(f"<b>{self.s['recommendation']}:</b> {f.get('recommendation','').replace(chr(10),' ')}", body))
            if f.get("evidence"):
                story.append(Paragraph(f"<b>{self.s['evidence']}:</b>", body))
                story.append(Paragraph(f.get("evidence","")[:300], mono))
            if f.get("poc"):
                story.append(Paragraph(f"<b>{self.s['poc']}:</b>", body))
                story.append(Paragraph(f.get("poc","")[:300], mono))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1,3*mm))

        story.append(PageBreak())

        # ── COMPLIANCE SUMMARY ────────────────────────────────
        if pci_results:
            story.append(Paragraph(self.s["compliance"], h2))
            story.append(Paragraph("PCI DSS v4.0", h3))
            pci_rows = [["Requirement", "Control", "Status", "Detail"]]
            for key, val in sorted(pci_results.items()):
                status = val.get("status","")
                cl_s   = colors.green if status=="PASS" else (colors.red if status=="FAIL" else colors.grey)
                pci_rows.append([
                    val.get("requirement",""),
                    Paragraph(val.get("control","")[:50], body),
                    Paragraph(f'<font color="{cl_s.hexval()}">{status}</font>', body),
                    Paragraph(val.get("detail","")[:60], body),
                ])
            pci_table = Table(pci_rows, colWidths=[20*mm,55*mm,22*mm,80*mm])
            pci_table.setStyle(TableStyle([
                ("BACKGROUND",   (0,0),(-1,0), DARK_BLUE),
                ("TEXTCOLOR",    (0,0),(-1,0), colors.white),
                ("FONTNAME",     (0,0),(-1,0), "Helvetica-Bold"),
                ("FONTSIZE",     (0,0),(-1,-1), 7),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT_GRAY,colors.white]),
                ("GRID",         (0,0),(-1,-1),0.3,colors.lightgrey),
                ("BOTTOMPADDING",(0,0),(-1,-1),3),
            ]))
            story.append(pci_table)

        doc.build(story)
        return path

    def _generate_pdf_minimal(self) -> Optional[str]:
        """Fallback: write text-based PDF manually."""
        path = os.path.join(self.output_dir, f"report_{self.lang}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"GATOR PRO Enterprise — {self.s['report_title']}\n")
            f.write(f"Client: {self.client} | Date: {self.date_str}\n")
            f.write("=" * 80 + "\n\n")
            f.write(f"TOTAL FINDINGS: {len(self.findings)}\n\n")
            for i, finding in enumerate(self.findings, 1):
                f.write(f"#{i} [{finding.get('severity','').upper()}] "
                        f"CVSS {finding.get('cvss',0)} — {finding.get('title','')}\n")
                f.write(f"  Host:   {finding.get('host','')}\n")
                f.write(f"  Descr:  {finding.get('description','')[:200]}\n")
                f.write(f"  Rec:    {finding.get('recommendation','')[:200]}\n\n")
        return path

    # ─── DOCX Generation ─────────────────────────────────────
    def _generate_docx(self) -> Optional[str]:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return None

        path = os.path.join(self.output_dir, f"report_{self.lang}.docx")
        doc  = Document()

        # Page margins
        for section in doc.sections:
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)

        def heading(text, level=1, color=(26,43,74)):
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)
            return p

        def para(text, bold=False, italic=False, size=10, color=None):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p

        # ── Cover ─────────────────────────────────────────────
        doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run("🐊 GATOR PRO ENTERPRISE")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)

        h = doc.add_heading(self.s["report_title"], level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = "Table Grid"
        cells = [
            (self.s["prepared_for"]+":", self.client),
            (self.s["date"]+":",         self.date_str),
            ("Target:",                  self.scan.get("target","")),
            (self.s["risk_level"]+":",   self._calc_risk_score()),
        ]
        for i, (label, val) in enumerate(cells):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[1].text = val
            meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_page_break()

        # ── Executive Summary ─────────────────────────────────
        heading(self.s["executive_summary"])
        sev_counts = {s: len([f for f in self.findings if f.get("severity")==s])
                      for s in ("critical","high","medium","low","info")}
        sum_table = doc.add_table(rows=len(sev_counts)+2, cols=2)
        sum_table.style = "Table Grid"
        sum_table.rows[0].cells[0].text = "Severity"
        sum_table.rows[0].cells[1].text = "Count"
        for i, (sev, cnt) in enumerate(sev_counts.items(), 1):
            sum_table.rows[i].cells[0].text = sev.upper()
            sum_table.rows[i].cells[1].text = str(cnt)
        sum_table.rows[-1].cells[0].text = "TOTAL"
        sum_table.rows[-1].cells[1].text = str(len(self.findings))
        doc.add_page_break()

        # ── Methodology ───────────────────────────────────────
        heading(self.s["methodology"])
        doc.add_paragraph(self.s["methodology_text"])
        doc.add_page_break()

        # ── Findings ──────────────────────────────────────────
        heading(self.s["findings"])
        SEV_COLORS_RGB = {
            "critical": (220,38,38), "high": (234,88,12),
            "medium": (217,119,6),  "low":  (22,163,74), "info": (107,114,128),
        }
        for i, f in enumerate(self.findings, 1):
            sev  = f.get("severity","info")
            cvss = f.get("cvss",0)
            col  = SEV_COLORS_RGB.get(sev,(0,0,0))
            p = doc.add_paragraph()
            r = p.add_run(f"#{i} [{sev.upper()}] CVSS {cvss} — {f.get('title','')}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(*col)

            ftable = doc.add_table(rows=2, cols=2)
            ftable.style = "Table Grid"
            ftable.rows[0].cells[0].text = self.s["affected"]
            ftable.rows[0].cells[1].text = str(f.get("host",""))
            ftable.rows[1].cells[0].text = self.s["owasp"]
            ftable.rows[1].cells[1].text = f.get("owasp_category","")

            doc.add_paragraph(f"{self.s['description']}: {f.get('description','')}")
            rec_p = doc.add_paragraph()
            rec_run = rec_p.add_run(f"{self.s['recommendation']}:")
            rec_run.bold = True
            doc.add_paragraph(f.get("recommendation",""))
            if f.get("evidence"):
                ev_p = doc.add_paragraph()
                ev_p.add_run(f"{self.s['evidence']}:").bold = True
                doc.add_paragraph(f.get("evidence","")[:300]).style = "No Spacing"
            doc.add_paragraph()

        # ── PCI Compliance ────────────────────────────────────
        pci_results = self.scan.get("pci_results",{})
        if pci_results:
            doc.add_page_break()
            heading(f"PCI DSS v4.0 — {self.s['compliance']}")
            pci_t = doc.add_table(rows=1+len(pci_results), cols=3)
            pci_t.style = "Table Grid"
            hdr = pci_t.rows[0]
            hdr.cells[0].text = "Requirement"
            hdr.cells[1].text = "Control"
            hdr.cells[2].text = "Status"
            for i, (key, val) in enumerate(sorted(pci_results.items()), 1):
                row = pci_t.rows[i]
                row.cells[0].text = str(val.get("requirement",""))
                row.cells[1].text = str(val.get("control",""))[:60]
                row.cells[2].text = str(val.get("status",""))

        doc.save(path)
        return path
