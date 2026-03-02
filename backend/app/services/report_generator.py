"""
GATOR PRO Enterprise — Report Generator
═════════════════════════════════════════
Generates professional pentest reports:
  • PDF — executive summary + detailed technical report (RU/EN)
  • DOCX — editable Word document for editing/branding
  • Telegram — instant findings notification via Bot API
  • JSON — machine-readable structured output

Report structure:
  1. Cover page (target, date, assessor, classification)
  2. Executive Summary (risk rating, top-10 findings, compliance)
  3. Scope and Methodology
  4. Findings by severity (Critical → Info)
     - Title, CVSS score, OWASP category
     - Description (business impact)
     - Technical evidence + PoC
     - Remediation steps with code examples
     - PCI DSS / SWIFT control mapping
  5. PCI DSS v4.0 Compliance Table
  6. SWIFT CSP Control Verification Matrix
  7. Appendix — Tool output, raw data
"""

import json
import os
import re
import time
import urllib.request
import urllib.parse
import ssl
from datetime import datetime
from pathlib import Path
from typing import Optional


# ─── CVSS → Risk label ────────────────────────────────────────
def cvss_risk(cvss: float) -> str:
    if cvss >= 9.0:  return "CRITICAL"
    if cvss >= 7.0:  return "HIGH"
    if cvss >= 4.0:  return "MEDIUM"
    if cvss > 0:     return "LOW"
    return "INFO"

def cvss_color_hex(sev: str) -> str:
    return {"critical":"#DC143C","high":"#FF6B35","medium":"#FFB300",
            "low":"#4CAF50","info":"#2196F3"}.get(sev.lower(),"#9E9E9E")

def risk_emoji(sev: str) -> str:
    return {"critical":"🔴","high":"🟠","medium":"🟡","low":"🟢","info":"🔵"}.get(sev.lower(),"⚪")


class ReportGenerator:
    """Full report generator: PDF, DOCX, JSON, Telegram."""

    def __init__(self, scan_data: dict, output_dir: str = None):
        self.scan    = scan_data
        self.target  = scan_data.get("target","unknown")
        self.scan_id = scan_data.get("id","")
        # Cross-platform: use /app/reports/output inside Docker, or system temp
        if output_dir is None:
            import tempfile
            output_dir = os.environ.get("REPORTS_DIR",
                        os.path.join(os.path.dirname(__file__), "..", "..", "reports", "output"))
        self.out_dir = Path(output_dir)
        self.out_dir.mkdir(parents=True, exist_ok=True)
        self.ts      = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        self.findings = self._collect_all_findings()
        self._sort_findings()

    def _collect_all_findings(self) -> list:
        """Aggregate findings from all module results."""
        all_f = []
        results = self.scan.get("raw_results") or {}
        if isinstance(results, str):
            try:
                results = json.loads(results)
            except Exception:
                results = {}

        # Collect from all module outputs
        for module_key in results:
            module_data = results[module_key]
            if isinstance(module_data, dict):
                findings = module_data.get("findings", [])
                if isinstance(findings, list):
                    all_f.extend(findings)

        # Also collect from scan.findings relation (if passed directly)
        direct_findings = self.scan.get("findings",[])
        if isinstance(direct_findings, list):
            all_f.extend(direct_findings)

        # Deduplicate by title+host
        seen = set()
        unique = []
        for f in all_f:
            key = f"{f.get('host','')}:{f.get('title','')[:60]}"
            if key not in seen:
                seen.add(key)
                unique.append(f)
        return unique

    def _sort_findings(self):
        sev_order = {"critical":0,"high":1,"medium":2,"low":3,"info":4}
        self.findings.sort(
            key=lambda f: (sev_order.get(f.get("severity","info"),5), -f.get("cvss",0)))

    # ═══════════════════════════════════════════════════════════
    def generate_all(self, telegram_token: str = None,
                     telegram_chat_id: str = None) -> dict:
        """Generate all report formats. Returns paths."""
        paths = {}

        # Always generate JSON
        paths["json"] = self.generate_json()

        # PDF
        try:
            paths["pdf"] = self.generate_pdf()
        except Exception as e:
            paths["pdf_error"] = str(e)

        # DOCX
        try:
            paths["docx"] = self.generate_docx()
        except Exception as e:
            paths["docx_error"] = str(e)

        # Telegram
        if telegram_token and telegram_chat_id:
            try:
                self.send_telegram(telegram_token, telegram_chat_id)
                paths["telegram"] = "sent"
            except Exception as e:
                paths["telegram_error"] = str(e)

        return paths

    # ─── JSON Report ──────────────────────────────────────────
    def generate_json(self) -> str:
        fname = self.out_dir / f"report_{self.target}_{self.ts}.json"
        report = {
            "report_version":  "2.0",
            "generated_at":    datetime.utcnow().isoformat(),
            "target":          self.target,
            "scan_id":         self.scan_id,
            "summary":         self._build_summary(),
            "findings":        self.findings,
            "compliance":      self._build_compliance_table(),
            "scan_data":       self.scan.get("raw_results",{}),
        }
        with open(fname, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=2, ensure_ascii=False, default=str)
        return str(fname)

    def _build_summary(self) -> dict:
        counts = {"critical":0,"high":0,"medium":0,"low":0,"info":0}
        for f in self.findings:
            sev = f.get("severity","info").lower()
            counts[sev] = counts.get(sev,0) + 1
        total = sum(counts.values())
        # Risk score: weighted sum
        risk_score = (
            counts["critical"] * 10 +
            counts["high"]     *  7 +
            counts["medium"]   *  4 +
            counts["low"]      *  1
        )
        risk_level = "CRITICAL" if counts["critical"] > 0 else \
                     "HIGH"     if counts["high"]     > 0 else \
                     "MEDIUM"   if counts["medium"]   > 0 else \
                     "LOW"      if counts["low"]      > 0 else "INFO"
        return {
            "total_findings":  total,
            "by_severity":     counts,
            "risk_score":      risk_score,
            "risk_level":      risk_level,
            "top_cve_count":   len([f for f in self.findings if f.get("cve_ids")]),
            "pci_fail_count":  len([f for f in self.findings if f.get("pci_dss_req")]),
            "swift_fail_count":len([f for f in self.findings if f.get("swift_control")]),
        }

    def _build_compliance_table(self) -> dict:
        results = self.scan.get("raw_results",{})
        if isinstance(results, str):
            try: results = json.loads(results)
            except: results = {}
        pci_data = results.get("pci_swift",{})
        if isinstance(pci_data, dict):
            return pci_data.get("summary",{})
        return {}

    # ─── PDF Report ───────────────────────────────────────────
    def generate_pdf(self) -> str:
        """Generate PDF using ReportLab (installed in Docker) or HTML fallback."""
        fname = str(self.out_dir / f"report_{self.target}_{self.ts}.pdf")
        try:
            return self._generate_pdf_reportlab(fname)
        except ImportError:
            # Fallback: generate HTML and convert with weasyprint or wkhtmltopdf
            return self._generate_pdf_html_fallback(fname)

    def _generate_pdf_reportlab(self, fname: str) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, PageBreak, HRFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

        doc = SimpleDocTemplate(
            fname, pagesize=A4,
            leftMargin=20*mm, rightMargin=20*mm,
            topMargin=25*mm, bottomMargin=20*mm,
        )
        styles = getSampleStyleSheet()
        story  = []

        # Color palette
        DARK_BG    = colors.HexColor("#1a1a2e")
        ACCENT     = colors.HexColor("#e94560")
        LIGHT_GREY = colors.HexColor("#f4f4f4")
        TEXT_COLOR = colors.HexColor("#2c2c2c")

        # Custom styles
        title_style = ParagraphStyle("title",
            fontSize=28, fontName="Helvetica-Bold",
            textColor=ACCENT, alignment=TA_CENTER, spaceAfter=10)
        sub_style = ParagraphStyle("sub",
            fontSize=14, fontName="Helvetica",
            textColor=TEXT_COLOR, alignment=TA_CENTER, spaceAfter=6)
        h1_style = ParagraphStyle("h1",
            fontSize=16, fontName="Helvetica-Bold",
            textColor=DARK_BG, spaceBefore=14, spaceAfter=8)
        body_style = ParagraphStyle("body",
            fontSize=9, fontName="Helvetica",
            textColor=TEXT_COLOR, leading=14, alignment=TA_JUSTIFY)
        code_style = ParagraphStyle("code",
            fontSize=8, fontName="Courier",
            textColor=colors.HexColor("#006400"),
            backColor=colors.HexColor("#f0f8f0"),
            leftIndent=10, rightIndent=10, leading=12)

        summary = self._build_summary()
        now     = datetime.utcnow().strftime("%d.%m.%Y %H:%M UTC")
        scan_type = self.scan.get("scan_type","Full Pentest").upper()

        # ── Cover Page ──────────────────────────────────────
        story.append(Spacer(1, 30*mm))
        story.append(Paragraph("🐊 GATOR PRO Enterprise", title_style))
        story.append(Paragraph("Banking Security Assessment Report", sub_style))
        story.append(Spacer(1, 10*mm))
        story.append(HRFlowable(width="100%", color=ACCENT, thickness=2))
        story.append(Spacer(1, 8*mm))

        cover_data = [
            ["Target:",     self.target],
            ["Scan Type:",  scan_type],
            ["Date:",       now],
            ["Scan ID:",    str(self.scan_id)[:16] + "..."],
            ["Risk Level:", summary["risk_level"]],
        ]
        cover_table = Table(cover_data, colWidths=[50*mm, 120*mm])
        cover_table.setStyle(TableStyle([
            ("FONTNAME",  (0,0),(-1,-1), "Helvetica"),
            ("FONTSIZE",  (0,0),(-1,-1), 11),
            ("FONTNAME",  (0,0),(0,-1),  "Helvetica-Bold"),
            ("TEXTCOLOR", (0,0),(0,-1),  DARK_BG),
            ("TOPPADDING",(0,0),(-1,-1), 4),
            ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ]))
        story.append(cover_table)
        story.append(Spacer(1, 15*mm))

        # Risk level badge
        risk_colors = {
            "CRITICAL": "#DC143C","HIGH":"#FF6B35",
            "MEDIUM":"#FFB300","LOW":"#4CAF50","INFO":"#2196F3"
        }
        risk_bg = colors.HexColor(risk_colors.get(summary["risk_level"],"#9E9E9E"))
        risk_style = ParagraphStyle("risk",
            fontSize=20, fontName="Helvetica-Bold",
            textColor=colors.white, backColor=risk_bg,
            alignment=TA_CENTER, borderPadding=8)
        story.append(Paragraph(
            f"OVERALL RISK: {summary['risk_level']}", risk_style))
        story.append(PageBreak())

        # ── Executive Summary ────────────────────────────────
        story.append(Paragraph("1. Executive Summary", h1_style))
        story.append(HRFlowable(width="100%", color=ACCENT, thickness=1))
        story.append(Spacer(1, 4*mm))

        exec_text = (
            f"This report presents the results of an automated security assessment "
            f"performed against <b>{self.target}</b> on {now}. "
            f"A total of <b>{summary['total_findings']}</b> security findings were identified."
        )
        story.append(Paragraph(exec_text, body_style))
        story.append(Spacer(1, 6*mm))

        # Findings summary table
        sev_data = [["Severity","Count","Risk Weight","Action Required"]]
        for sev, weight, action in [
            ("CRITICAL", counts := summary["by_severity"]["critical"],
             "×10", "Immediate — within 24 hours"),
            ("HIGH",     summary["by_severity"]["high"],
             "×7",  "Urgent — within 7 days"),
            ("MEDIUM",   summary["by_severity"]["medium"],
             "×4",  "Normal — within 30 days"),
            ("LOW",      summary["by_severity"]["low"],
             "×1",  "Planned — within 90 days"),
            ("INFO",     summary["by_severity"]["info"],
             "×0",  "Advisory — at discretion"),
        ]:
            sev_data.append([sev, str(weight) if isinstance(weight,int) else "0", weight if isinstance(weight,str) else "×0", action])

        # Fix: rebuild properly
        sev_rows = [["Severity","Count","Weight","Action Required"]]
        for sev_name, count_key, weight_str, action_str in [
            ("🔴 CRITICAL", "critical", "×10", "Within 24 hours"),
            ("🟠 HIGH",     "high",     "×7",  "Within 7 days"),
            ("🟡 MEDIUM",   "medium",   "×4",  "Within 30 days"),
            ("🟢 LOW",      "low",      "×1",  "Within 90 days"),
            ("🔵 INFO",     "info",     "×0",  "Advisory"),
        ]:
            cnt = summary["by_severity"].get(count_key, 0)
            sev_rows.append([sev_name, str(cnt), weight_str, action_str])

        sev_table = Table(sev_rows, colWidths=[45*mm, 20*mm, 20*mm, 85*mm])
        sev_table.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0), DARK_BG),
            ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
            ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT_GREY, colors.white]),
            ("GRID",       (0,0),(-1,-1), 0.5, colors.grey),
            ("ALIGN",      (1,0),(2,-1), "CENTER"),
            ("TOPPADDING", (0,0),(-1,-1), 4),
            ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ]))
        story.append(sev_table)
        story.append(Spacer(1, 6*mm))

        # PCI/SWIFT summary
        story.append(Paragraph(
            f"<b>PCI DSS Violations:</b> {summary['pci_fail_count']} | "
            f"<b>SWIFT CSP Failures:</b> {summary['swift_fail_count']} | "
            f"<b>CVE References:</b> {summary['top_cve_count']}",
            body_style))
        story.append(PageBreak())

        # ── Detailed Findings ────────────────────────────────
        story.append(Paragraph("2. Technical Findings", h1_style))
        story.append(HRFlowable(width="100%", color=ACCENT, thickness=1))
        story.append(Spacer(1, 4*mm))

        for i, finding in enumerate(self.findings, 1):
            sev  = finding.get("severity","info").lower()
            cvss = finding.get("cvss", 0)
            title= finding.get("title","")
            sev_bg = colors.HexColor(cvss_color_hex(sev))

            # Finding header
            find_style = ParagraphStyle(f"find_title_{i}",
                fontSize=10, fontName="Helvetica-Bold",
                textColor=sev_bg, spaceBefore=8, spaceAfter=4)
            story.append(Paragraph(
                f"{i}. [{sev.upper()} | CVSS {cvss}] {title}", find_style))

            # Metadata row
            meta = []
            if finding.get("owasp_category"):
                meta.append(f"OWASP: {finding['owasp_category']}")
            if finding.get("pci_dss_req"):
                meta.append(f"PCI DSS: {', '.join(finding['pci_dss_req'])}")
            if finding.get("swift_control"):
                meta.append(f"SWIFT: {', '.join(finding['swift_control'])}")
            if finding.get("cve_ids"):
                meta.append(f"CVE: {', '.join(finding['cve_ids'][:3])}")
            if meta:
                story.append(Paragraph(" | ".join(meta),
                    ParagraphStyle("meta", fontSize=8, textColor=colors.grey,
                                   spaceAfter=3, fontName="Helvetica-Oblique")))

            # Description
            desc = finding.get("description","")
            if desc:
                story.append(Paragraph(f"<b>Description:</b> {desc[:500]}", body_style))
                story.append(Spacer(1, 2*mm))

            # Evidence
            evidence = finding.get("evidence","")
            if evidence:
                story.append(Paragraph("<b>Evidence:</b>", body_style))
                story.append(Paragraph(evidence[:400], code_style))
                story.append(Spacer(1, 2*mm))

            # Recommendation
            rec = finding.get("recommendation","")
            if rec:
                story.append(Paragraph(f"<b>Remediation:</b> {rec[:400]}", body_style))

            story.append(HRFlowable(width="80%", color=LIGHT_GREY, thickness=0.5))

            # Page break every 4 findings
            if i % 4 == 0:
                story.append(PageBreak())

        story.append(PageBreak())

        # ── PCI DSS Compliance Table ─────────────────────────
        story.append(Paragraph("3. PCI DSS v4.0 Compliance Summary", h1_style))
        story.append(HRFlowable(width="100%", color=ACCENT, thickness=1))
        story.append(Spacer(1, 4*mm))

        results = self.scan.get("raw_results",{})
        if isinstance(results, str):
            try: results = json.loads(results)
            except: results = {}
        compliance_data = results.get("pci_swift",{})
        checks = []
        if isinstance(compliance_data, dict):
            checks = compliance_data.get("checks",[])

        if checks:
            pci_checks  = [c for c in checks if c.get("framework")=="PCI DSS"]
            swift_checks = [c for c in checks if c.get("framework")=="SWIFT CSP"]

            for framework_name, check_list in [
                ("PCI DSS v4.0", pci_checks),
                ("SWIFT CSP v2024", swift_checks),
            ]:
                if not check_list:
                    continue
                story.append(Paragraph(framework_name, ParagraphStyle("fw",
                    fontSize=12, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=4)))

                tbl_data = [["Req ID","Control","Status"]]
                for c in check_list:
                    status = c.get("status","")
                    tbl_data.append([
                        c.get("req_id",""),
                        Paragraph(c.get("title","")[:60], ParagraphStyle("td",
                            fontSize=8, fontName="Helvetica")),
                        status,
                    ])
                tbl = Table(tbl_data, colWidths=[20*mm, 120*mm, 25*mm])
                tbl_style = [
                    ("BACKGROUND",(0,0),(-1,0), DARK_BG),
                    ("TEXTCOLOR", (0,0),(-1,0), colors.white),
                    ("FONTNAME",  (0,0),(-1,0), "Helvetica-Bold"),
                    ("FONTSIZE",  (0,0),(-1,-1), 8),
                    ("GRID",      (0,0),(-1,-1), 0.3, colors.lightgrey),
                    ("TOPPADDING",(0,0),(-1,-1), 3),
                    ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                ]
                # Color status cells
                for row_idx, c in enumerate(check_list, 1):
                    status = c.get("status","")
                    bg = {"PASS":colors.HexColor("#e8f5e9"),
                          "FAIL":colors.HexColor("#ffebee"),
                          "WARNING":colors.HexColor("#fff8e1"),
                          "MANUAL":colors.HexColor("#e3f2fd")}.get(status, colors.white)
                    tbl_style.append(("BACKGROUND",(2,row_idx),(2,row_idx), bg))
                tbl.setStyle(TableStyle(tbl_style))
                story.append(tbl)
                story.append(Spacer(1, 4*mm))
        else:
            story.append(Paragraph("No compliance checks data available.", body_style))

        # Build PDF
        doc.build(story)
        return fname

    def _generate_pdf_html_fallback(self, fname: str) -> str:
        """Generate HTML → PDF using wkhtmltopdf or weasyprint."""
        html_content = self._generate_html_report()
        html_path = fname.replace(".pdf", "_temp.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        # Try wkhtmltopdf
        import subprocess
        try:
            result = subprocess.run(
                ["wkhtmltopdf","--quiet","--page-size","A4",html_path, fname],
                capture_output=True, timeout=60)
            if result.returncode == 0:
                os.unlink(html_path)
                return fname
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Try weasyprint
        try:
            import weasyprint
            weasyprint.HTML(filename=html_path).write_pdf(fname)
            os.unlink(html_path)
            return fname
        except ImportError:
            pass

        # Fallback: return HTML as "pdf"
        return html_path

    def _generate_html_report(self) -> str:
        """Generate styled HTML report."""
        summary = self._build_summary()
        now  = datetime.utcnow().strftime("%d.%m.%Y %H:%M UTC")
        rl   = summary["risk_level"]
        rl_color = {"CRITICAL":"#DC143C","HIGH":"#FF6B35","MEDIUM":"#FFB300",
                    "LOW":"#4CAF50","INFO":"#2196F3"}.get(rl,"#9E9E9E")

        findings_html = ""
        for i, f in enumerate(self.findings, 1):
            sev   = f.get("severity","info")
            color = cvss_color_hex(sev)
            pci   = ", ".join(f.get("pci_dss_req",[]))
            swift = ", ".join(f.get("swift_control",[]))
            cves  = ", ".join(f.get("cve_ids",[])[:3])
            findings_html += f"""
            <div class="finding">
              <div class="finding-header" style="border-left: 4px solid {color}">
                <span class="sev-badge" style="background:{color}">{sev.upper()}</span>
                <span class="cvss">CVSS {f.get("cvss",0)}</span>
                <strong>{i}. {f.get("title","")}</strong>
              </div>
              <div class="finding-body">
                {"<p><b>OWASP:</b> " + f.get("owasp_category","") + "</p>" if f.get("owasp_category") else ""}
                {"<p><b>PCI DSS:</b> " + pci + "</p>" if pci else ""}
                {"<p><b>SWIFT:</b> " + swift + "</p>" if swift else ""}
                {"<p><b>CVE:</b> " + cves + "</p>" if cves else ""}
                <p><b>Description:</b> {f.get("description","")[:400]}</p>
                {"<pre class='evidence'>" + f.get("evidence","")[:300] + "</pre>" if f.get("evidence") else ""}
                {"<p><b>Remediation:</b> " + f.get("recommendation","")[:300] + "</p>" if f.get("recommendation") else ""}
              </div>
            </div>"""

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>GATOR PRO — Security Report — {self.target}</title>
  <style>
    body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 0; padding: 20px; color: #2c2c2c; }}
    .cover {{ text-align: center; padding: 60px 0; background: linear-gradient(135deg,#1a1a2e,#16213e); color: white; margin: -20px -20px 30px; }}
    .cover h1 {{ font-size: 2.5em; color: #e94560; margin: 0; }}
    .cover h2 {{ font-size: 1.3em; font-weight: normal; color: #aaa; }}
    .risk-badge {{ display: inline-block; padding: 12px 40px; font-size: 1.8em; font-weight: bold; color: white; background: {rl_color}; border-radius: 6px; margin: 20px 0; }}
    .meta {{ background: #f4f4f4; padding: 15px; border-radius: 6px; margin: 20px 0; }}
    .meta table {{ width: 100%; border-collapse: collapse; }}
    .meta td {{ padding: 6px 12px; }}
    .meta td:first-child {{ font-weight: bold; width: 150px; color: #1a1a2e; }}
    h2 {{ color: #1a1a2e; border-bottom: 2px solid #e94560; padding-bottom: 8px; }}
    .summary-grid {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 10px; margin: 20px 0; }}
    .sev-card {{ text-align: center; padding: 15px; border-radius: 8px; color: white; font-weight: bold; }}
    .finding {{ margin: 15px 0; border: 1px solid #e0e0e0; border-radius: 6px; overflow: hidden; page-break-inside: avoid; }}
    .finding-header {{ padding: 10px 15px; background: #f9f9f9; display: flex; align-items: center; gap: 10px; }}
    .sev-badge {{ padding: 3px 10px; border-radius: 3px; color: white; font-size: 0.8em; font-weight: bold; }}
    .cvss {{ font-size: 0.85em; color: #666; }}
    .finding-body {{ padding: 12px 15px; }}
    .finding-body p {{ margin: 6px 0; font-size: 0.9em; }}
    pre.evidence {{ background: #f0f8f0; border: 1px solid #c8e6c9; padding: 10px; font-size: 0.8em; overflow-x: auto; border-radius: 4px; }}
    @media print {{ .finding {{ page-break-inside: avoid; }} }}
  </style>
</head>
<body>
<div class="cover">
  <h1>🐊 GATOR PRO Enterprise</h1>
  <h2>Banking Security Assessment Report</h2>
  <div class="risk-badge">{rl}</div>
  <p style="color:#ccc; font-size:0.9em">Generated: {now}</p>
</div>

<div class="meta">
  <table>
    <tr><td>Target:</td><td><strong>{self.target}</strong></td></tr>
    <tr><td>Scan ID:</td><td>{str(self.scan_id)[:20]}</td></tr>
    <tr><td>Date:</td><td>{now}</td></tr>
    <tr><td>Total Findings:</td><td><strong>{summary["total_findings"]}</strong></td></tr>
    <tr><td>PCI DSS Failures:</td><td>{summary["pci_fail_count"]}</td></tr>
    <tr><td>SWIFT CSP Failures:</td><td>{summary["swift_fail_count"]}</td></tr>
  </table>
</div>

<h2>Executive Summary</h2>
<div class="summary-grid">
  <div class="sev-card" style="background:#DC143C">🔴 CRITICAL<br><span style="font-size:2em">{summary["by_severity"]["critical"]}</span></div>
  <div class="sev-card" style="background:#FF6B35">🟠 HIGH<br><span style="font-size:2em">{summary["by_severity"]["high"]}</span></div>
  <div class="sev-card" style="background:#FFB300">🟡 MEDIUM<br><span style="font-size:2em">{summary["by_severity"]["medium"]}</span></div>
  <div class="sev-card" style="background:#4CAF50">🟢 LOW<br><span style="font-size:2em">{summary["by_severity"]["low"]}</span></div>
  <div class="sev-card" style="background:#2196F3">🔵 INFO<br><span style="font-size:2em">{summary["by_severity"]["info"]}</span></div>
</div>

<h2>Detailed Findings</h2>
{findings_html}

<hr>
<p style="text-align:center; color:#999; font-size:0.8em">
  GATOR PRO Enterprise v2.0 — Confidential — {now}
</p>
</body>
</html>"""

    # ─── DOCX Report ──────────────────────────────────────────
    def generate_docx(self) -> str:
        fname = str(self.out_dir / f"report_{self.target}_{self.ts}.docx")
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        summary = self._build_summary()
        now = datetime.utcnow().strftime("%d.%m.%Y %H:%M UTC")

        def add_heading(text, level=1, color=(26,26,46)):
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.color.rgb = RGBColor(*color)
            return h

        def add_colored_para(text, color=(44,44,44), bold=False, size=10):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(*color)
            run.font.bold = bold
            run.font.size = Pt(size)
            return p

        # Title
        add_heading("🐊 GATOR PRO Enterprise", level=1)
        add_heading("Banking Security Assessment Report", level=2, color=(100,100,100))
        doc.add_paragraph()

        # Cover info table
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"
        cover_data = [
            ("Target:", self.target),
            ("Date:", now),
            ("Risk Level:", summary["risk_level"]),
            ("Total Findings:", str(summary["total_findings"])),
            ("PCI DSS Failures:", str(summary["pci_fail_count"])),
        ]
        for i, (label, value) in enumerate(cover_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True

        doc.add_page_break()

        # Executive Summary
        add_heading("1. Executive Summary")
        doc.add_paragraph(
            f"This report presents automated security assessment results for "
            f"{self.target} ({now}). "
            f"Total findings: {summary['total_findings']}."
        )

        # Severity table
        add_heading("Findings by Severity", level=2)
        stbl = doc.add_table(rows=6, cols=4)
        stbl.style = "Medium Grid 1 Accent 1"
        headers = ["Severity","Count","CVSS Weight","Action"]
        for j, h in enumerate(headers):
            stbl.cell(0, j).text = h
            stbl.cell(0, j).paragraphs[0].runs[0].bold = True
        rows = [
            ("CRITICAL", str(summary["by_severity"]["critical"]), "×10", "24 hours"),
            ("HIGH",     str(summary["by_severity"]["high"]),     "×7",  "7 days"),
            ("MEDIUM",   str(summary["by_severity"]["medium"]),   "×4",  "30 days"),
            ("LOW",      str(summary["by_severity"]["low"]),      "×1",  "90 days"),
            ("INFO",     str(summary["by_severity"]["info"]),     "×0",  "Advisory"),
        ]
        for i, row in enumerate(rows, 1):
            for j, val in enumerate(row):
                stbl.cell(i, j).text = val

        doc.add_page_break()

        # Detailed Findings
        add_heading("2. Technical Findings")
        for i, f in enumerate(self.findings, 1):
            sev   = f.get("severity","info").upper()
            cvss  = f.get("cvss", 0)
            title = f.get("title","")
            color_map = {
                "CRITICAL":(220,20,60),"HIGH":(255,107,53),
                "MEDIUM":(255,179,0),"LOW":(76,175,80),"INFO":(33,150,243)
            }
            c = color_map.get(sev,(100,100,100))
            add_heading(f"{i}. {title}", level=3, color=c)
            meta_parts = [f"Severity: {sev}  |  CVSS: {cvss}"]
            if f.get("owasp_category"):
                meta_parts.append(f"OWASP: {f['owasp_category']}")
            if f.get("pci_dss_req"):
                meta_parts.append(f"PCI DSS: {', '.join(f['pci_dss_req'])}")
            doc.add_paragraph("  |  ".join(meta_parts)).italic = True

            if f.get("description"):
                p = doc.add_paragraph()
                p.add_run("Description: ").bold = True
                p.add_run(f.get("description","")[:400])

            if f.get("evidence"):
                p = doc.add_paragraph()
                p.add_run("Evidence: ").bold = True
                p.add_run(f.get("evidence","")[:300])

            if f.get("recommendation"):
                p = doc.add_paragraph()
                p.add_run("Remediation: ").bold = True
                p.add_run(f.get("recommendation","")[:400])

            doc.add_paragraph()

        doc.add_page_break()

        # PCI compliance
        add_heading("3. PCI DSS v4.0 Compliance")
        results = self.scan.get("raw_results",{})
        if isinstance(results, str):
            try: results = json.loads(results)
            except: results = {}
        compliance = results.get("pci_swift",{})
        checks = compliance.get("checks",[]) if isinstance(compliance, dict) else []
        pci_checks = [c for c in checks if c.get("framework")=="PCI DSS"]

        if pci_checks:
            ptbl = doc.add_table(rows=1+len(pci_checks), cols=3)
            ptbl.style = "Light Grid Accent 1"
            for j, h in enumerate(["Req ID","Control","Status"]):
                ptbl.cell(0, j).text = h
                ptbl.cell(0, j).paragraphs[0].runs[0].bold = True
            for i, c in enumerate(pci_checks, 1):
                ptbl.cell(i, 0).text = c.get("req_id","")
                ptbl.cell(i, 1).text = c.get("title","")[:70]
                ptbl.cell(i, 2).text = c.get("status","")

        doc.save(fname)
        return fname

    # ─── Telegram ────────────────────────────────────────────
    def send_telegram(self, token: str, chat_id: str):
        """Send findings summary to Telegram bot."""
        summary = self._build_summary()
        now = datetime.utcnow().strftime("%d.%m.%Y %H:%M")

        rl = summary["risk_level"]
        rl_emoji = {"CRITICAL":"🔴","HIGH":"🟠","MEDIUM":"🟡","LOW":"🟢","INFO":"🔵"}.get(rl,"⚪")

        # Main summary message
        msg = (
            f"🐊 *GATOR PRO — Scan Complete*\n\n"
            f"🎯 Target: `{self.target}`\n"
            f"📅 Date: {now}\n"
            f"⚠️ Risk: {rl_emoji} *{rl}*\n\n"
            f"📊 *Findings:*\n"
            f"🔴 Critical: {summary['by_severity']['critical']}\n"
            f"🟠 High:     {summary['by_severity']['high']}\n"
            f"🟡 Medium:   {summary['by_severity']['medium']}\n"
            f"🟢 Low:      {summary['by_severity']['low']}\n"
            f"🔵 Info:     {summary['by_severity']['info']}\n\n"
            f"📋 PCI DSS Failures: {summary['pci_fail_count']}\n"
            f"🏦 SWIFT CSP Failures: {summary['swift_fail_count']}"
        )
        self._tg_send(token, chat_id, msg)

        # Critical/High findings
        criticals = [f for f in self.findings
                     if f.get("severity","").lower() in ("critical","high")][:5]
        if criticals:
            details = "🚨 *Top Critical/High Findings:*\n\n"
            for i, f in enumerate(criticals, 1):
                sev = f.get("severity","").upper()
                emoji = risk_emoji(f.get("severity","info"))
                details += (
                    f"{emoji} *{i}. [{sev} | CVSS {f.get('cvss',0)}]*\n"
                    f"`{f.get('title','')[:60]}`\n"
                )
                if f.get("url"):
                    details += f"URL: {f.get('url','')[:60]}\n"
                details += "\n"
            self._tg_send(token, chat_id, details)

    def _tg_send(self, token: str, chat_id: str, text: str):
        """Send Telegram message via Bot API."""
        url  = f"https://api.telegram.org/bot{token}/sendMessage"
        data = urllib.parse.urlencode({
            "chat_id":    chat_id,
            "text":       text[:4096],
            "parse_mode": "Markdown",
        }).encode()
        req = urllib.request.Request(url, data=data, method="POST")
        req.add_header("Content-Type","application/x-www-form-urlencoded")
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        try:
            with urllib.request.urlopen(req, context=ctx, timeout=15) as resp:
                result = json.loads(resp.read())
                if not result.get("ok"):
                    raise ValueError(f"Telegram API error: {result}")
        except Exception as e:
            raise RuntimeError(f"Telegram send failed: {e}")


# ─── Celery task integration ─────────────────────────────────
def generate_report(scan_data: dict, output_dir: str,
                    telegram_token: str = None,
                    telegram_chat_id: str = None) -> dict:
    """Main entry point called from Celery task."""
    gen = ReportGenerator(scan_data, output_dir)
    return gen.generate_all(telegram_token, telegram_chat_id)
